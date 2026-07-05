from io import BytesIO

import pytest

from adapters.export.docx import export_docx


def _build_template_bytes() -> bytes:
    """A real .docx with a custom style, a header, and non-default margins —
    everything the template path must preserve while replacing the body."""
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches

    tpl = Document()
    tpl.styles.add_style("MioStileCustom", WD_STYLE_TYPE.PARAGRAPH)

    section = tpl.sections[0]
    section.header.paragraphs[0].text = "INTESTAZIONE AZIENDALE"
    section.left_margin = Inches(2)
    section.top_margin = Inches(2)

    tpl.add_paragraph("TESTO ORIGINALE DEL TEMPLATE")

    buf = BytesIO()
    tpl.save(buf)
    return buf.getvalue()


_DOCUMENT = {
    "title": "Nuovo Titolo",
    "content": {
        "type": "doc",
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "Contenuto generato"}]},
        ],
    },
}


def _read(out: bytes):
    from docx import Document

    return Document(BytesIO(out))


class TestExportDocxWithTemplate:
    def test_body_content_replaced(self):
        out = export_docx(_DOCUMENT, template_bytes=_build_template_bytes())
        text = "\n".join(p.text for p in _read(out).paragraphs)
        assert "TESTO ORIGINALE DEL TEMPLATE" not in text
        assert "Contenuto generato" in text

    def test_header_preserved(self):
        out = export_docx(_DOCUMENT, template_bytes=_build_template_bytes())
        header = _read(out).sections[0].header.paragraphs[0].text
        assert "INTESTAZIONE AZIENDALE" in header

    def test_margins_preserved(self):
        from docx.shared import Inches

        out = export_docx(_DOCUMENT, template_bytes=_build_template_bytes())
        section = _read(out).sections[0]
        assert section.left_margin == Inches(2)
        assert section.top_margin == Inches(2)

    def test_corrupt_bytes_raise_valueerror(self):
        with pytest.raises(ValueError, match="Template DOCX"):
            export_docx(_DOCUMENT, template_bytes=b"non e' un docx")


class TestExportDocxWithoutTemplate:
    def test_plain_export_still_works(self):
        out = export_docx(_DOCUMENT)
        text = "\n".join(p.text for p in _read(out).paragraphs)
        assert "Nuovo Titolo" in text
        assert "Contenuto generato" in text

    def test_missing_heading_style_falls_back(self):
        """A template stripped of the 'Heading 1' style must not raise — the
        renderer degrades to a bold paragraph instead of a KeyError."""
        from docx import Document

        tpl = Document()
        style = tpl.styles["Heading 1"]
        style.element.getparent().remove(style.element)
        buf = BytesIO()
        tpl.save(buf)

        out = export_docx(_DOCUMENT, template_bytes=buf.getvalue())
        text = "\n".join(p.text for p in _read(out).paragraphs)
        assert "Nuovo Titolo" in text
