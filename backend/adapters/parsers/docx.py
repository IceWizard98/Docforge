from dataclasses import dataclass, field


@dataclass
class DocxSection:
    heading: str = ""
    content: str = ""
    style: str = ""


@dataclass
class DocxParsedDocument:
    sections: list[DocxSection] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


def _parse_docx_paragraphs(doc) -> list[DocxSection]:
    sections: list[DocxSection] = []
    heading = ""
    content_lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            if content_lines or heading:
                sections.append(
                    DocxSection(
                        heading=heading,
                        content="\n".join(content_lines),
                        style="paragraph",
                    )
                )
                content_lines = []
            heading = text
        else:
            content_lines.append(text)
    if content_lines or heading:
        sections.append(
            DocxSection(
                heading=heading,
                content="\n".join(content_lines),
                style="paragraph",
            )
        )
    return sections


def _parse_docx_tables(doc) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        table_data: list[list[str]] = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables


def parse_docx(file_path: str) -> DocxParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    doc = DocxDocument(file_path)
    return DocxParsedDocument(
        sections=_parse_docx_paragraphs(doc),
        tables=_parse_docx_tables(doc),
        metadata={"file": file_path},
    )


def parse_docx_bytes(data: bytes) -> DocxParsedDocument:
    try:
        from io import BytesIO

        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    doc = DocxDocument(BytesIO(data))
    return DocxParsedDocument(
        sections=_parse_docx_paragraphs(doc),
        tables=_parse_docx_tables(doc),
    )
