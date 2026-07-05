from contextlib import suppress
from io import BytesIO

from adapters.export._prosemirror import walk


def _plain(runs: list[tuple[str, list[str]]]) -> str:
    return "".join(text for text, _ in runs)


def _safe_heading(doc, text: str, level: int):
    # A custom template may lack the built-in 'Heading N' styles; degrade to a
    # bold paragraph rather than raising KeyError.
    try:
        return doc.add_heading(text, level=min(level, 9))
    except KeyError:
        para = doc.add_paragraph()
        para.add_run(text).bold = True
        return para


class _DocxRenderer:
    def __init__(self, doc) -> None:
        self.doc = doc

    def _styled_paragraph(self, style: str):
        # Template may not define the list/named style — fall back to plain.
        try:
            return self.doc.add_paragraph(style=style)
        except KeyError:
            return self.doc.add_paragraph()

    def _add_inline(self, para, runs) -> None:
        from docx.shared import Pt
        for text, marks in runs:
            run = para.add_run(text)
            run.font.size = Pt(11)
            if "bold" in marks:
                run.bold = True
            if "italic" in marks:
                run.italic = True
            if "code" in marks:
                run.font.name = "Courier New"
                run.font.size = Pt(10)

    def heading(self, level: int, runs) -> None:
        _safe_heading(self.doc, _plain(runs), level)

    def paragraph(self, runs) -> None:
        self._add_inline(self.doc.add_paragraph(), runs)

    def bullet_list(self, items) -> None:
        for runs in items:
            self._add_inline(self._styled_paragraph("List Bullet"), runs)

    def ordered_list(self, items) -> None:
        for runs in items:
            self._add_inline(self._styled_paragraph("List Number"), runs)

    def blockquote(self, kids) -> None:
        from adapters.export._prosemirror import inline_runs
        for p in kids:
            para = self.doc.add_paragraph()
            self._add_inline(para, [("> ", ["bold"])])
            self._add_inline(para, inline_runs(p.get("content", [])))

    def code_block(self, lang: str, code: str) -> None:
        from docx.shared import Pt
        run = self.doc.add_paragraph().add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    def horizontal_rule(self) -> None:
        self.doc.add_paragraph("─" * 60)

    def table(self, rows) -> None:
        if not rows:
            return
        cols = max((len(cells) for cells in rows), default=0)
        if cols == 0:
            return
        table = self.doc.add_table(rows=len(rows), cols=cols)
        # Template without the grid style: leave the default table style.
        with suppress(KeyError):
            table.style = "Table Grid"
        for row_idx, cells in enumerate(rows):
            for col_idx, (_is_header, runs) in enumerate(cells):
                if col_idx < cols:
                    table.cell(row_idx, col_idx).text = _plain(runs)


def export_docx(document: dict, template_bytes: bytes | None = None) -> bytes:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    if template_bytes is not None:
        try:
            doc = DocxDocument(BytesIO(template_bytes))
        except Exception as exc:
            raise ValueError(f"Template DOCX non valido: {exc}") from exc
        # Clear the template body but keep the trailing sectPr — it carries page
        # size, margins and the header/footer references, which must survive.
        body = doc.element.body
        for el in list(body):
            if not el.tag.endswith("}sectPr"):
                body.remove(el)
    else:
        doc = DocxDocument()

    _safe_heading(doc, document.get("title", "Document"), 1)

    content = document.get("content", {})
    if isinstance(content, dict):
        walk(content.get("content", []), _DocxRenderer(doc))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
